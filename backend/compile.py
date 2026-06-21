"""Compile included chapters into a standard manuscript-format .docx.

Follows the common "Shunn modern" novel conventions: 1" margins, 12pt Times
New Roman, double-spaced, first-line indent, a title page with contact info and
an approximate word count, a running header (Surname / TITLE / page) from page
two, chapters starting on new pages, and centered "#" scene breaks.
"""
import io
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "blockquote", "li"}


class _HTMLBlocks(HTMLParser):
    """Turn a chapter's HTML into a list of blocks of styled runs."""
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self.cur = None
        self.btype = "p"
        self.b = self.i = self.u = 0

    def _ensure(self):
        if self.cur is None:
            self.cur = {"type": self.btype, "runs": []}

    def flush(self):
        if self.cur is not None:
            if any(t.strip() for t, *_ in self.cur["runs"]):
                self.blocks.append(self.cur)
            self.cur = None
        self.btype = "p"

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in BLOCK_TAGS:
            self.flush()
            self.btype = "h" if tag in ("h1", "h2", "h3") else ("quote" if tag == "blockquote" else "p")
            self._ensure()
        elif tag == "br":
            self._ensure()
            self.cur["runs"].append(("\n", self.b > 0, self.i > 0, self.u > 0))
        elif tag == "hr":
            self.flush()
            self.blocks.append({"type": "scene", "runs": [("#", False, False, False)]})
        elif tag in ("b", "strong"):
            self.b += 1
        elif tag in ("i", "em"):
            self.i += 1
        elif tag == "u":
            self.u += 1

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in BLOCK_TAGS:
            self.flush()
        elif tag in ("b", "strong"):
            self.b = max(0, self.b - 1)
        elif tag in ("i", "em"):
            self.i = max(0, self.i - 1)
        elif tag == "u":
            self.u = max(0, self.u - 1)

    def handle_data(self, data):
        if not data:
            return
        self._ensure()
        self.cur["runs"].append((data, self.b > 0, self.i > 0, self.u > 0))


def _parse(html: str):
    p = _HTMLBlocks()
    p.feed(html or "")
    p.close()
    p.flush()
    # Plain-text fallback: if no block tags were present, split on blank lines.
    if not p.blocks and (html or "").strip():
        return [{"type": "p", "runs": [(line, False, False, False)]}
                for line in html.split("\n") if line.strip()]
    return p.blocks


def _double(par):
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _single(par):
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)


def _page_field(par):
    run = par.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def build_manuscript(title: str, author: str, sheets) -> bytes:
    """sheets: ordered list of objects with .title, .text (HTML), .words."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.different_first_page_header_footer = True

    surname = author.split()[-1] if author.strip() else "Author"
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.add_run(f"{surname} / {(title or 'UNTITLED').upper()} / ")
    _page_field(hp)

    # ---- title page ----
    total = sum(s.words for s in sheets)
    rounded = round(total / 1000) * 1000 if total >= 1000 else round(total / 100) * 100

    contact = doc.add_paragraph(); _single(contact)
    contact.add_run(author or "Author Name").add_break()
    contact.add_run("Contact address").add_break()
    contact.add_run("email@example.com")

    wc = doc.add_paragraph(); wc.alignment = WD_ALIGN_PARAGRAPH.RIGHT; _single(wc)
    wc.add_run(f"About {rounded:,} words")

    for _ in range(8):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run((title or "Untitled").upper()).bold = True
    by = doc.add_paragraph(); by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by.add_run("by " + (author or "Author Name"))

    # ---- chapters ----
    for idx, s in enumerate(sheets):
        doc.add_page_break()
        head = doc.add_paragraph(); head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.add_run(s.title or f"Chapter {idx + 1}").bold = True
        doc.add_paragraph()  # blank line before prose

        for blk in _parse(s.text or ""):
            p = doc.add_paragraph(); _double(p)
            if blk["type"] == "scene":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("#")
                continue
            if blk["type"] == "h":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif blk["type"] == "p":
                p.paragraph_format.first_line_indent = Inches(0.5)
            for (text, b, i, u) in blk["runs"]:
                for j, seg in enumerate(text.split("\n")):
                    if j > 0:
                        p.add_run().add_break()
                    if seg:
                        r = p.add_run(seg)
                        r.bold, r.italic, r.underline = b, i, u

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
